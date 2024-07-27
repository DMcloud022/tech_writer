import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx2pdf import convert
import logging
import matplotlib.pyplot as plt
import io
from concurrent.futures import ThreadPoolExecutor
import re
from typing import List, Tuple

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.max_workers = 4  # Adjust based on your system's capabilities

    def format_document(self, content: str, metadata: dict = None) -> Document:
        doc = Document()
        self.set_document_styles(doc)

        try:
            self.add_title_page(doc, metadata)
            self.add_table_of_contents(doc)
            self.add_content(doc, content)
            self.add_page_number(doc.sections[0])
            self.update_table_of_contents(doc)
        except Exception as e:
            logger.error(f"Error formatting document: {e}")
            raise

        return doc

    def add_title_page(self, doc: Document, metadata: dict = None):
        title = metadata.get('title', 'Generated Technical Document') if metadata else 'Generated Technical Document'
        author = metadata.get('author', 'Technical Writing Assistant') if metadata else 'Technical Writing Assistant'
        date = metadata.get('date', pd.Timestamp.now().strftime('%Y-%m-%d')) if metadata else pd.Timestamp.now().strftime('%Y-%m-%d')

        doc.add_paragraph(title, style='Title')
        doc.add_paragraph(f"Prepared by: {author}", style='Subtitle')
        doc.add_paragraph(f"Date: {date}", style='Subtitle')
        doc.add_page_break()

    def add_table_of_contents(self, doc: Document):
        doc.add_paragraph("Table of Contents", style='Heading 1')
        doc.add_paragraph("", style='Normal')
        doc.add_page_break()

    def add_content(self, doc: Document, content: str):
        if not content:
            raise ValueError("Content is empty")
        
        sections = self.parse_content(content)
        
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            futures = []
            for section_type, section_content in sections:
                futures.append(executor.submit(self.process_section, doc, section_type, section_content))
            
            for future in futures:
                future.result()

    def parse_content(self, content: str) -> List[Tuple[str, str]]:
        sections = []
        current_type = 'Normal'
        current_content = []

        for line in content.split('\n'):
            stripped_line = line.strip()
            if stripped_line.startswith('# '):
                if current_content:
                    sections.append((current_type, '\n'.join(current_content)))
                current_type = 'Heading 1'
                current_content = [stripped_line.lstrip('# ')]
            elif stripped_line.startswith('## '):
                if current_content:
                    sections.append((current_type, '\n'.join(current_content)))
                current_type = 'Heading 2'
                current_content = [stripped_line.lstrip('## ')]
            elif stripped_line.startswith('### '):
                if current_content:
                    sections.append((current_type, '\n'.join(current_content)))
                current_type = 'Heading 3'
                current_content = [stripped_line.lstrip('### ')]
            elif '[CHART]' in stripped_line:
                if current_content:
                    sections.append((current_type, '\n'.join(current_content)))
                sections.append(('Chart', stripped_line))
                current_type = 'Normal'
                current_content = []
            elif '[TABLE]' in stripped_line:
                if current_content:
                    sections.append((current_type, '\n'.join(current_content)))
                sections.append(('Table', stripped_line))
                current_type = 'Normal'
                current_content = []
            elif '[IMAGE]' in stripped_line:
                if current_content:
                    sections.append((current_type, '\n'.join(current_content)))
                sections.append(('Image', stripped_line))
                current_type = 'Normal'
                current_content = []
            else:
                current_content.append(stripped_line)

        if current_content:
            sections.append((current_type, '\n'.join(current_content)))

        return sections

    def process_section(self, doc: Document, section_type: str, section_content: str):
        if section_type.startswith('Heading'):
            doc.add_paragraph(section_content, style=section_type)
        elif section_type == 'Chart':
            self.add_chart_placeholder(doc, section_content)
        elif section_type == 'Table':
            self.add_table_placeholder(doc, section_content)
        elif section_type == 'Image':
            self.add_image_placeholder(doc, section_content)
        else:
            doc.add_paragraph(section_content, style='Normal')

    def set_document_styles(self, doc: Document):
        styles = doc.styles

        style_specs = {
            'Title': {'font_size': Pt(24), 'bold': True, 'alignment': WD_ALIGN_PARAGRAPH.CENTER},
            'Subtitle': {'font_size': Pt(18), 'italic': True, 'color': RGBColor(100, 100, 100), 'alignment': WD_ALIGN_PARAGRAPH.CENTER},
            'Heading 1': {'font_size': Pt(16), 'bold': True, 'space_after': Pt(12)},
            'Heading 2': {'font_size': Pt(14), 'bold': True, 'space_after': Pt(6)},
            'Heading 3': {'font_size': Pt(12), 'bold': True, 'space_after': Pt(6)},
            'Normal': {'font_size': Pt(11), 'line_spacing_rule': WD_LINE_SPACING.SINGLE, 'space_after': Pt(6)},
        }

        for style_name, specs in style_specs.items():
            self.create_or_update_style(styles, style_name, **specs)

    def create_or_update_style(self, styles, style_name: str, **kwargs):
        style = styles[style_name] if style_name in styles else styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

        font = style.font
        font.name = 'Calibri'
        font.size = kwargs.get('font_size', Pt(11))
        font.bold = kwargs.get('bold', False)
        font.italic = kwargs.get('italic', False)
        if 'color' in kwargs:
            font.color.rgb = kwargs['color']

        paragraph_format = style.paragraph_format
        if 'alignment' in kwargs:
            paragraph_format.alignment = kwargs['alignment']
        if 'space_after' in kwargs:
            paragraph_format.space_after = kwargs['space_after']
        if 'line_spacing_rule' in kwargs:
            paragraph_format.line_spacing_rule = kwargs['line_spacing_rule']

    def add_page_number(self, section):
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.text = "Page "
        self.add_field_code(paragraph, "PAGE")
        paragraph.add_run(" of ")
        self.add_field_code(paragraph, "NUMPAGES")

    def add_field_code(self, paragraph, field_code: str):
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = field_code
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.extend([fldChar1, instrText, fldChar2])

    def add_chart_placeholder(self, doc: Document, content: str):
        doc.add_paragraph("Chart Placeholder", style='Normal')
        chart_type = re.search(r'\[CHART:(\w+)\]', content)
        chart_type = chart_type.group(1) if chart_type else 'generic'
        self.add_placeholder_image(doc, f'{chart_type}_chart.png')

    def add_table_placeholder(self, doc: Document, content: str):
        doc.add_paragraph("Table Placeholder", style='Normal')
        table_size = re.search(r'\[TABLE:(\d+)x(\d+)\]', content)
        rows, cols = (3, 3) if not table_size else map(int, table_size.groups())
        table = doc.add_table(rows=rows, cols=cols)
        for row in table.rows:
            for cell in row.cells:
                cell.text = "Data"

    def add_image_placeholder(self, doc: Document, content: str):
        doc.add_paragraph("Image Placeholder", style='Normal')
        image_type = re.search(r'\[IMAGE:(\w+)\]', content)
        image_type = image_type.group(1) if image_type else 'generic'
        self.add_placeholder_image(doc, f'{image_type}_image.png')

    def add_placeholder_image(self, doc: Document, filename: str):
        try:
            with io.BytesIO() as output:
                plt.figure(figsize=(4, 3))
                plt.text(0.5, 0.5, filename.split('.')[0], horizontalalignment='center', verticalalignment='center', fontsize=20, alpha=0.5)
                plt.axis('off')
                plt.savefig(output, format='png', dpi=300, bbox_inches='tight')
                plt.close()
                output.seek(0)
                doc.add_picture(output, width=Inches(4))
        except Exception as e:
            logger.error(f"Failed to add placeholder image: {e}")

    def update_table_of_contents(self, doc: Document):
        toc_p = None
        for paragraph in doc.paragraphs:
            if paragraph.style.name == 'TOC Heading':
                toc_p = paragraph
                break
        if toc_p:
            toc_p._element.getparent().remove(toc_p._element)
        doc.add_paragraph("", style='Normal')
        self.add_field_code(doc.paragraphs[-1], "TOC \\o \"1-3\" \\h \\z \\u")

    def save_document(self, doc: Document, file_path: str):
        try:
            _, ext = os.path.splitext(file_path)
            ext = ext.lower()

            if ext == '.docx':
                doc.save(file_path)
                logger.info(f"Document saved as DOCX: {file_path}")
            elif ext == '.pdf':
                docx_path = file_path.replace('.pdf', '.docx')
                doc.save(docx_path)
                logger.info(f"Document saved as DOCX: {docx_path}")
                self.convert_docx_to_pdf(docx_path, file_path)
            else:
                logger.error(f"Unsupported file extension: {ext}")
                raise ValueError(f"Unsupported file extension: {ext}")
        except Exception as e:
            logger.error(f"Failed to save document: {e}")
            raise

    def convert_docx_to_pdf(self, docx_path: str, pdf_path: str):
        try:
            convert(docx_path, pdf_path)
            logger.info(f"Converted DOCX to PDF: {pdf_path}")
        except Exception as e:
            logger.error(f"Failed to convert DOCX to PDF: {e}")
            logger.error("Ensure 'docx2pdf' is installed and properly configured.")
            raise